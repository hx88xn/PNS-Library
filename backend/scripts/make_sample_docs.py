#!/usr/bin/env python
"""Generate a sample corpus for development and for the refusal eval.

These are synthetic documents written in the register of a design office, not
real Pakistan Navy material. They exist so the pipeline can be exercised
end to end — every parser, every collection — before real documents are
available on the server.

    python scripts/make_sample_docs.py [--out sample-docs]
"""

from __future__ import annotations

import argparse
from pathlib import Path

import ezdxf
import fitz
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook

# ── Source material ──────────────────────────────────────────────────────
# (section heading, body) pairs per document.

STABILITY = [
    (
        "4.2 Righting arm curve",
        "For frigate and corvette hull forms the area under the righting arm curve up to "
        "30 degrees shall not be less than 0.055 m·rad, and the area between 30 and 40 "
        "degrees not less than 0.030 m·rad. Maximum GZ shall occur at a heel angle of not "
        "less than 30 degrees, with a value of at least 0.20 m. Initial transverse "
        "metacentric height GM in the full load departure condition shall be not less than "
        "0.15 m after correction for free surface effect in all partially filled tanks.",
    ),
    (
        "4.5 Free surface correction",
        "Free surface moments shall be computed for every tank capable of being partially "
        "filled in service, taken at the filling level producing the greatest moment. Tanks "
        "declared as either pressed full or empty in a given condition may be excluded only "
        "where a positive means of maintaining that state is fitted.",
    ),
    (
        "6.1 Extent of damage",
        "The platform shall survive flooding of any two adjacent main watertight "
        "compartments. Assumed damage extent is 15 percent of waterline length or 21 m, "
        "whichever is less, penetrating to the centreline and extending vertically without "
        "limit above the waterline. In the final equilibrium condition the margin line shall "
        "not be submerged, residual GZ shall be positive over a range of at least 15 "
        "degrees, and the angle of list shall not exceed 20 degrees.",
    ),
    (
        "6.4 Permeability",
        "Permeability shall be taken as 95 percent for void spaces and accommodation, 85 "
        "percent for machinery spaces, and 60 percent for spaces occupied by stores. Where "
        "a compartment's permeability is materially different, it shall be calculated and "
        "recorded in the damage stability booklet.",
    ),
]

SEAKEEPING = [
    (
        "4.1 Motion limits",
        "Helicopter launch and recovery shall be possible up to sea state 5 with significant "
        "wave height of 4.0 m. Limiting motions at the flight deck reference point are 5 "
        "degrees single amplitude roll, 3 degrees pitch, and 2.0 m per second significant "
        "vertical velocity. Where passive tank or fin stabilisation is fitted, the criteria "
        "shall be met with the stabilisation system operating and the residual case shall be "
        "reported separately.",
    ),
    (
        "5.2 Active fin selection",
        "Active fin stabilisers shall reduce significant roll amplitude by not less than 70 "
        "percent at the design speed in beam seas at the natural roll period. Fin area shall "
        "be sized on the roll damping requirement at 18 knots; performance falls with the "
        "square of ship speed and the fins shall not be relied upon below 8 knots. Fin "
        "housings shall not intrude into a main machinery space or breach a watertight "
        "boundary below the damage control deck.",
    ),
    (
        "2.3 Length to beam ratio",
        "A length to beam ratio between 7.5 and 8.5 is recommended for patrol and escort "
        "platforms in the 2,000 to 3,500 tonne displacement band, balancing resistance in "
        "the design speed range against transverse stability margin and deck area for combat "
        "systems. Block coefficient in the range 0.44 to 0.50 has proven satisfactory for a "
        "design speed of 28 knots.",
    ),
]

STRUCTURE = [
    (
        "5.1 Still water bending moment",
        "The hull girder section modulus at amidships shall be calculated for both hogging "
        "and sagging conditions, combining the still water bending moment from the worst "
        "realistic loading condition with the wave bending moment derived from the applicable "
        "classification rule. A minimum margin of 10 percent above the rule-required section "
        "modulus shall be retained at the initial design stage to absorb weight growth over "
        "the service life.",
    ),
    (
        "5.4 Shell plating",
        "Transverse frame spacing shall be 600 mm throughout the machinery spaces and 700 mm "
        "elsewhere unless the arrangement requires otherwise. Bottom shell plating thickness "
        "in way of the machinery spaces shall not be less than 9 mm in grade AH36 steel. "
        "Where plating thickness changes between adjacent strakes, the difference shall not "
        "exceed 3 mm and the taper shall extend over not less than three frame spaces.",
    ),
    (
        "9.2 Shock grades",
        "Equipment essential to propulsion, steering, firefighting and combat capability "
        "shall be qualified to Grade A shock. Grade B applies to equipment whose failure "
        "would not immediately impair fighting capability but could injure personnel or "
        "damage Grade A items. Deck-mounted Grade A equipment shall be qualified by hammer "
        "or barge test, and mounting seats shall be shown to transmit no amplification "
        "greater than 1.5 in the 5 to 100 Hz band.",
    ),
]

SURVIVABILITY = [
    (
        "2.1 Citadel and zones",
        "The platform shall be divided into not fewer than four damage control zones, each "
        "self-sufficient in firefighting, dewatering and electrical supply for a period of 30 "
        "minutes without support from adjacent zones. Zone boundaries shall coincide with "
        "main transverse watertight bulkheads and shall carry A-60 fire insulation. The fire "
        "main shall be a ring with sectioning valves at each zone boundary so that any single "
        "rupture isolates no more than one zone.",
    ),
    (
        "4.3 Superstructure faceting",
        "Superstructure sides shall be inclined not less than 7 degrees from the vertical, "
        "with the inclination held constant over each face to avoid creating a specular "
        "return band. Vertical corner reflectors formed by intersecting deckhouse faces and "
        "deck edges shall be eliminated by chamfering or by screening. Deck fittings above "
        "the damage control deck shall be either recessed, screened, or shaped to avoid "
        "dihedral geometry.",
    ),
    (
        "5.2 Exhaust cooling",
        "Gas turbine exhaust shall be diluted with ambient air to reduce plume temperature "
        "below 200 degrees Celsius at the funnel exit. Uptake and funnel surfaces reaching "
        "temperatures above 60 degrees Celsius shall be insulated or air-gap screened from "
        "external view. The funnel shall be positioned so that the plume does not wash the "
        "flight deck approach path in the prevailing relative wind envelope.",
    ),
]

PROPULSION = [
    (
        "3.1 Plant configuration",
        "The combined diesel and gas arrangement pairs two cruise diesel engines with a "
        "single boost gas turbine through a cross-connect gearbox. Cruise diesels drive both "
        "shafts at speeds up to 18 knots; the gas turbine is clutched in for sprint operation "
        "above that. The cross-connect gearbox is the single most space-critical item in the "
        "machinery arrangement and its footprint shall be fixed before the main transverse "
        "bulkhead positions are frozen.",
    ),
    (
        "4.5 Blade loading",
        "Five-bladed controllable pitch propellers shall be adopted to reduce blade rate "
        "excitation into the hull. Cavitation inception speed shall be not less than 15 knots "
        "at the design draught to preserve the acoustic advantage during transit. Blade area "
        "ratio shall be selected so that mean thrust loading does not exceed 130 kN per "
        "square metre of expanded blade area at continuous service rating.",
    ),
    (
        "2.2 Alignment calculation",
        "Shaft alignment shall be calculated for the hot running condition with the hull in "
        "the full load afloat state. No bearing shall be unloaded in any operating condition, "
        "and the reaction on the aftmost sterntube bearing shall remain between 60 and 110 "
        "percent of its nominal design load.",
    ),
]

STAFF_REQUIREMENT = [
    (
        "1.1 Mission profile",
        "The platform shall sustain a 21-day patrol without replenishment, with an endurance "
        "of 4,500 nautical miles at 14 knots economical speed. Maximum continuous speed shall "
        "be not less than 24 knots in sea state 3 at full load displacement. The platform "
        "shall operate a 10-tonne class helicopter and shall launch and recover a 7 m rigid "
        "hull inflatable boat at speeds up to 10 knots in sea state 4.",
    ),
    (
        "2.4 Accommodation standard",
        "Design complement is 96 with margin for a further 20 embarked personnel. Junior "
        "rates shall be berthed in cabins of not more than six, with a minimum of 2.8 square "
        "metres of deck area and 0.85 cubic metres of stowage per person. No accommodation "
        "space shall have its only escape route through a machinery space.",
    ),
    (
        "6.1 Weight margin policy",
        "A design and build margin of 6 percent of lightship weight and a KG margin of 0.25 m "
        "shall be carried through to contract design. Service life allowance shall be a "
        "further 5 percent of lightship displacement and 0.15 m of KG, reserved against "
        "equipment growth over 30 years.",
    ),
]

STANDARDS = [
    (
        "Part 3, Clause 12",
        "Watertight doors below the damage control deck shall be of the quick-acting clip "
        "type, operable from both sides, and shall be proof tested to a head of water equal "
        "to the distance from the sill to the bulkhead deck plus 1 m. Weathertight hatches on "
        "exposed decks shall be fitted with coamings of not less than 380 mm on the freeboard "
        "deck and 230 mm on superstructure decks.",
    ),
    (
        "Annex D",
        "Where an aluminium superstructure meets a steel hull, the joint shall be made with "
        "an explosion-bonded bimetallic transition strip. Direct bolted connections between "
        "aluminium and steel are not permitted. The transition strip shall be protected from "
        "standing water on both faces.",
    ),
    (
        "3.3 Impressed current system",
        "An impressed current cathodic protection system shall maintain hull potential "
        "between -0.80 and -1.05 V against a silver/silver chloride reference cell. Anode "
        "positions shall be kept clear of the propeller aperture and of sonar transducer "
        "faces by not less than 3 m. The underwater coating scheme shall be a glass flake "
        "epoxy system of 400 micrometre nominal dry film thickness.",
    ),
]


def write_docx(path: Path, doc_ref: str, title: str, revision: str, sections) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    document.add_heading(title, level=0)
    header = document.add_paragraph()
    header.add_run(f"{doc_ref}   {revision}   RESTRICTED").bold = True

    for heading, body in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(body)

    document.save(str(path))


def write_pdf(path: Path, doc_ref: str, title: str, revision: str, sections) -> None:
    """One section per page, so page-number citations are exercised end to end.

    A citation that says p.55 has to survive being checked against the sheet,
    and a single-page fixture would never catch a parser that loses the page.
    """
    document = fitz.open()

    cover = document.new_page()
    cover.insert_text((72, 96), title, fontsize=18, fontname="hebo")
    cover.insert_text((72, 124), f"{doc_ref}   {revision}", fontsize=10, fontname="helv")
    cover.insert_text((72, 142), "RESTRICTED", fontsize=10, fontname="hebo")

    for heading, body in sections:
        page = document.new_page()
        page.insert_text((72, 72), heading, fontsize=13, fontname="hebo")
        page.insert_textbox(
            fitz.Rect(72, 96, 523, 700), body, fontsize=10, fontname="helv", align=0
        )
        # Running footer, as a controlled document would carry.
        page.insert_text(
            (72, 760), f"{doc_ref}   {revision}   RESTRICTED", fontsize=8, fontname="helv"
        )

    document.save(str(path))
    document.close()


def write_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Loading conditions"
    sheet.append(["Condition", "Displacement (t)", "KG (m)", "GM fluid (m)", "Draught (m)", "Trim (m)"])
    for row in [
        ("Full load departure", 3120, 5.42, 0.86, 4.15, 0.12),
        ("Full load arrival", 2870, 5.51, 0.74, 3.98, -0.08),
        ("Half consumables", 2995, 5.47, 0.80, 4.06, 0.02),
        ("Light seagoing", 2640, 5.58, 0.66, 3.81, -0.21),
        ("Lightship", 2385, 5.63, 0.59, 3.62, -0.34),
    ]:
        sheet.append(row)

    tanks = workbook.create_sheet("Tank capacities")
    tanks.append(["Tank", "Contents", "Capacity (m3)", "Frame from", "Frame to", "Free surface moment (t·m)"])
    for row in [
        ("F1 P/S", "Diesel fuel", 148.0, 42, 54, 61.2),
        ("F2 P/S", "Diesel fuel", 132.5, 54, 66, 55.8),
        ("FW1", "Fresh water", 64.0, 70, 78, 18.4),
        ("BW1 P/S", "Ballast", 96.5, 24, 34, 40.1),
    ]:
        tanks.append(row)

    workbook.save(str(path))


def write_dxf(path: Path) -> None:
    document = ezdxf.new("R2010", setup=True)
    modelspace = document.modelspace()

    block = document.blocks.new(name="TITLEBLOCK")
    block.add_attdef("DWG_NO", (0, 0), height=2.5)
    block.add_attdef("TITLE", (0, -5), height=2.5)
    block.add_attdef("REV", (0, -10), height=2.5)
    block.add_attdef("SCALE", (0, -15), height=2.5)
    block.add_attdef("DRAWN", (0, -20), height=2.5)

    insert = modelspace.add_blockref("TITLEBLOCK", (0, 0))
    insert.add_auto_attribs(
        {
            "DWG_NO": "SDO/DRG/GA-1042",
            "TITLE": "Transverse watertight bulkhead, frame 62",
            "REV": "Rev B",
            "SCALE": "1:50",
            "DRAWN": "SDO Structures",
        }
    )

    for layer, texts in {
        "STRUCTURE": [
            "Bulkhead plating 8 mm AH36",
            "Vertical stiffeners 140 x 8 FB at 600 crs",
            "Stiffener brackets to deck 3 and deck 4",
        ],
        "NOTES": [
            "Bulkhead to be watertight to deck 3",
            "All welds continuous fillet 5 mm unless noted",
            "Penetrations to be fitted with approved WT glands",
        ],
        "PIPING": [
            "Fire main penetration at frame 62 port",
            "Bilge suction penetration below deck 5",
        ],
    }.items():
        document.layers.add(name=layer)
        for index, text in enumerate(texts):
            modelspace.add_text(
                text, dxfattribs={"layer": layer, "height": 2.5}
            ).set_placement((20, -index * 6 - 40))

    document.saveas(str(path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, default=Path(__file__).parent.parent / "sample-docs")
    args = parser.parse_args()

    out = args.out
    out.mkdir(parents=True, exist_ok=True)

    write_docx(
        out / "SDO-NA-STAB-014 Stability criteria.docx",
        "SDO/NA/STAB-014", "Stability criteria for surface combatants", "Rev C", STABILITY,
    )
    write_docx(
        out / "SDO-SUR-DC-018 Survivability.docx",
        "SDO/SUR/DC-018", "Survivability and signature management", "Rev B", SURVIVABILITY,
    )
    write_docx(
        out / "SDO-NSR-PATROL-2031 Staff requirement.docx",
        "SDO/NSR/PATROL-2031", "Staff requirement, offshore patrol vessel", "Rev F", STAFF_REQUIREMENT,
    )
    write_pdf(
        out / "SDO-STR-SCANT-021 Structural design.pdf",
        "SDO/STR/SCANT-021", "Structural design and scantlings", "Rev B", STRUCTURE,
    )
    write_pdf(
        out / "SDO-PROP-CODAG-009 Propulsion.pdf",
        "SDO/PROP/CODAG-009", "Propulsion and machinery arrangement", "Rev A", PROPULSION,
    )
    write_pdf(
        out / "SDO-NA-SEAKEEP-011 Seakeeping.pdf",
        "SDO/NA/SEAKEEP-011", "Seakeeping and hull form", "Rev B", SEAKEEPING,
    )
    write_pdf(
        out / "NES-109 Closure standards.pdf",
        "NES-109", "Watertight and weathertight closures", "Issue 4", STANDARDS,
    )
    write_xlsx(out / "SDO-NA-HYDRO-003 Loading conditions.xlsx")
    write_dxf(out / "SDO-DRG-GA-1042 Bulkhead frame 62.dxf")

    print(f"Wrote {len(list(out.iterdir()))} sample documents to {out}")


if __name__ == "__main__":
    main()
