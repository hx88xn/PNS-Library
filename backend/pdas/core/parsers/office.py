"""Word and Excel.

DOCX heading styles give real section boundaries — better structure than a PDF,
where headings have to be inferred from font size.

XLSX has no prose. Each row is rendered as `header: value` pairs so a retrieved
cell carries the meaning of its column, and the locator is the sheet and row.
A bare grid of numbers would embed to nothing useful.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook

from .base import Block, ParsedDocument, ParserError

_DOC_REF = re.compile(r"\b([A-Z]{2,}(?:[/-][A-Z0-9]{2,}){1,3})\b")
_REVISION = re.compile(r"\b(rev(?:ision)?\.?\s*[A-Z0-9]{1,3}|issue\s*\d+)\b", re.IGNORECASE)

MAX_ROWS = 5000
"""Guard against a spreadsheet with a million empty formatted rows."""


def parse_docx(path: Path) -> ParsedDocument:
    try:
        document = DocxDocument(str(path))
    except Exception as exc:
        raise ParserError(f"Cannot open DOCX: {exc}") from exc

    blocks: list[Block] = []
    current_section = ""

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        level = _heading_level(paragraph.style.name if paragraph.style else "")
        if level:
            current_section = text
            blocks.append(Block(text=text, section=current_section, heading_level=level))
        else:
            blocks.append(Block(text=text, section=current_section))

    # Tables carry the criteria in a lot of these documents; flatten each row.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                blocks.append(Block(text=line, section=current_section))

    head = " ".join(block.text for block in blocks[:12])[:2000]
    title = next((b.text for b in blocks if b.heading_level), None)

    return ParsedDocument(
        blocks=blocks,
        title=(title or path.stem).strip() or path.stem,
        doc_ref=_first(_DOC_REF, head),
        revision=_first(_REVISION, head),
        kind="text",
    )


def parse_xlsx(path: Path) -> ParsedDocument:
    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
    except Exception as exc:
        raise ParserError(f"Cannot open XLSX: {exc}") from exc

    blocks: list[Block] = []
    try:
        for sheet in workbook.worksheets:
            rows = sheet.iter_rows(values_only=True)
            headers: list[str] = []

            for index, row in enumerate(rows):
                if index > MAX_ROWS:
                    break
                values = ["" if cell is None else str(cell).strip() for cell in row]
                if not any(values):
                    continue

                if not headers:
                    headers = values
                    continue

                pairs = [
                    f"{headers[i] or f'col{i + 1}'}: {value}"
                    for i, value in enumerate(values)
                    if value and i < len(headers)
                ]
                if pairs:
                    blocks.append(
                        Block(
                            text="; ".join(pairs),
                            section=f"{sheet.title}, row {index + 1}",
                        )
                    )
    finally:
        workbook.close()

    return ParsedDocument(
        blocks=blocks,
        title=path.stem,
        doc_ref=_first(_DOC_REF, path.stem),
        kind="text",
    )


def _heading_level(style_name: str) -> int:
    if not style_name:
        return 0
    lowered = style_name.lower()
    if lowered.startswith("heading"):
        tail = lowered.replace("heading", "").strip()
        return int(tail) if tail.isdigit() else 1
    if lowered in {"title", "subtitle"}:
        return 1
    return 0


def _first(pattern: re.Pattern[str], text: str) -> str | None:
    match = pattern.search(text)
    return match.group(1).strip() if match else None
